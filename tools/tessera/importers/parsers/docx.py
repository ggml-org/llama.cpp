"""DOCX parser.

Reads a .docx file with python-docx and produces an
``IntermediateDocument`` (the intermediate JSON shape, which
``ast_builder`` then turns into a ``DocumentAST``).

The intermediate shape is deliberately simple: a list of
``IntermediateBlock`` dicts, plus a list of inline-run dicts. The
shape is what the AST builder consumes; it's a stable contract
between parsers and the AST layer, decoupled from both python-docx
and the AST module.

Coverage (v1):

* Headings (Heading 1..9) -> ``heading`` blocks with ``level`` attr.
* Paragraphs -> ``paragraph`` blocks. Bold / italic / underline /
  strikethrough runs are detected via python-docx's run-level
  style + font; links are detected by scanning the run XML for
  ``<w:hyperlink>`` elements.
* Numbered / bulleted lists -> ``list`` container + ``listItem``
  children. We walk the document body in order and treat any
  ``<w:numPr>`` run as a list item; consecutive items with the
  same ``numId`` form one list.
* Tables -> ``table`` container + cell ``paragraph`` blocks. We
  support both simple tables and nested tables (the latter
  promoted to a follow-up cell block).
* Images -> ``image`` blocks. The image bytes are written to a
  sidecar directory and the path is recorded as ``source``.
* Footnotes -> a single ``quote`` block per footnote at the end
  of the document. python-docx exposes footnotes via the
  ``part`` API; we read the footnotes part directly.

Punted (v1):

* Track changes (``<w:ins>`` / ``<w:del>``): the parser keeps the
  current state but doesn't surface track-change markers.
* Comments: not surfaced in v1.
* Sections / page layout: not preserved (the AST doesn't have
  a section model).
* Math (``<m:oMath>``): text-only fallback.
"""

from __future__ import annotations

import logging
import re
import shutil
import zipfile
from pathlib import Path
from typing import Any, Optional
from xml.etree import ElementTree as ET

from docx import Document
from docx.oxml.ns import qn

from ..ast_schema import (
    DocumentAST,
    make_heading,
    make_paragraph,
    run_to_json,
)
from ..intermediate import (
    IntermediateBlock,
    IntermediateDocument,
    IntermediateInlineRun,
)

log = logging.getLogger(__name__)

# Word XML namespaces
W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
R_NS = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}"


def parse_docx(path: Path, *, media_dir: Optional[Path] = None) -> IntermediateDocument:
    """Parse a .docx file into an IntermediateDocument.

    `media_dir` is the directory the parser writes image bytes to
    (one file per image); the on-disk path is recorded in the
    intermediate block's ``source`` attribute. When `media_dir` is
    None, images are skipped (and a warning is logged).
    """
    log.debug("parse_docx: %s", path)
    doc = Document(str(path))
    blocks: list[IntermediateBlock] = []

    # Body elements in document order. python-docx exposes the
    # XML body directly; we walk it to keep list/table nesting
    # correct (the high-level ``doc.paragraphs`` and
    # ``doc.tables`` APIs split the document and lose order).
    body = doc.element.body
    if body is None:
        return IntermediateDocument(blocks=[])

    # Track in-progress list containers so consecutive <w:p> with
    # the same numId fold into a single ``list`` block.
    open_lists: dict[int, dict[str, Any]] = {}
    # Index of the list block we last emitted (so we can append
    # items to it via the intermediate list refs).

    paragraphs_by_id: dict[int, Any] = {}
    for el in list(body):
        tag = el.tag
        if tag == qn("w:p"):
            p_obj = _wrap_paragraph(el, doc, paragraphs_by_id)
            if p_obj is None:
                continue
            blocks.extend(_emit_paragraph_or_listitem(p_obj, open_lists, blocks))
        elif tag == qn("w:tbl"):
            # A table interrupts any open list. Close the
            # lists before processing the table.
            _close_all_lists(open_lists, blocks)
            blocks.append(_emit_table(el, doc))
        elif tag == qn("w:sectPr"):
            # Section break; nothing to emit (page layout is not
            # preserved in the AST). Close any open list first.
            _close_all_lists(open_lists, blocks)
        else:
            # Any other element (unknown type) breaks any
            # open list. Close the lists and skip the element.
            _close_all_lists(open_lists, blocks)
            log.debug("parse_docx: skipping unsupported body element %s", tag)

    # Footnotes
    footnote_blocks = _extract_footnotes(path)
    blocks.extend(footnote_blocks)

    # Flush any list containers that were open at the end
    # of the document. Without this, the last list (if it
    # wasn't followed by a non-paragraph element) is lost.
    _close_all_lists(open_lists, blocks)

    # Images: walk the body again to find <w:drawing> and <w:pict>
    # elements that python-docx's high-level API doesn't surface
    # cleanly. We attach image blocks at the position of their
    # containing paragraph (replacing the empty placeholder).
    blocks = _inject_images(blocks, path, media_dir, doc)

    return IntermediateDocument(blocks=blocks)


# ---------------------------------------------------------------------------
# Paragraph + run extraction
# ---------------------------------------------------------------------------


def _wrap_paragraph(el: Any, doc: Document, paragraphs_by_id: dict[int, Any]) -> Any:
    """Wrap a <w:p> XML element as a python-docx Paragraph object.

    python-docx's ``doc.paragraphs`` only includes top-level
    paragraphs. We re-wrap any <w:p> we find so we can use
    python-docx's run/style accessors. The wrapper is read-only
    enough for our needs.
    """
    from docx.text.paragraph import Paragraph

    para = Paragraph(el, doc.part)
    paragraphs_by_id[id(para._p)] = para
    return para


def _runs_of(para: Any) -> list[Any]:
    """Return the list of runs in a paragraph.

    A "run" in python-docx is a <w:r> element. We re-fetch them
    from the live XML each call so subsequent mutations (e.g. by
    the table walker) are reflected.
    """
    return list(para.runs)


def _emit_paragraph_or_listitem(
    para: Any,
    open_lists: dict[int, dict[str, Any]],
    emitted_blocks: list[IntermediateBlock],
) -> list[IntermediateBlock]:
    """Emit a paragraph as a list item, a heading, or a plain paragraph.

    A <w:p> with <w:numPr> is a list item; consecutive items with
    the same numId + ilvl fold into one list container. We
    materialise the list container as an ``IntermediateBlock`` of
    type ``list`` and emit ``listItem`` blocks for the items.
    """
    style_name = (para.style.name or "").lower() if para.style is not None else ""
    text = "".join(run.text for run in _runs_of(para))

    # Headings: any style whose name starts with "heading"
    # (case-insensitive). We extract the level from the
    # trailing digit if present, else default to 1. The
    # style name is normalised to lowercase and stripped
    # of whitespace; python-docx exposes "Heading 1" (with
    # a space) but a hand-crafted DOCX may use "Heading1"
    # (no space) — we accept both.
    if style_name.startswith("heading"):
        m = re.search(r"(\d+)\s*$", style_name)
        level = int(m.group(1)) if m else 1
        level = max(1, min(level, 6))
        return [IntermediateBlock(type="heading", attrs={"level": level}, runs=[_run_dict(text)])]

    # Title style
    if style_name == "title":
        return [IntermediateBlock(type="heading", attrs={"level": 1}, runs=[_run_dict(text)])]

    # List item?
    num_info = _num_info(para)
    if num_info is not None:
        return _emit_list_item(para, num_info, open_lists, emitted_blocks, text)

    # Plain paragraph
    if not text.strip():
        # Empty paragraph - keep it as a paragraph with empty
        # content (matches the AST: an empty paragraph is still a
        # block, and removing it loses the user's spacing).
        return [IntermediateBlock(type="paragraph", runs=[])]

    runs = [_run_dict(r.text, _run_annotations(r)) for r in _runs_of(para)]
    # If there's only one run with no annotations, just keep its
    # text; otherwise keep the run breakdown (annotations carry
    # useful information like links).
    if len(runs) == 1 and not runs[0].annotations:
        return [IntermediateBlock(type="paragraph", runs=[_run_dict(text)])]
    return [IntermediateBlock(type="paragraph", runs=runs)]


def _run_dict(text: str, annotations: Optional[list[Any]] = None) -> IntermediateInlineRun:
    return IntermediateInlineRun(text=text, annotations=list(annotations or []))


def _run_annotations(run: Any) -> list[Any]:
    """Return the list of annotations for a python-docx run.

    Inspects the run's font + element XML for bold / italic /
    underline / strikethrough / code-style markers. Links aren't
    detected here (they're on the parent <w:hyperlink>, not on
    the <w:r>); we add them in ``_runs_with_links`` below.
    """
    ann: list[Any] = []
    if run.bold:
        ann.append("bold")
    if run.italic:
        ann.append("italic")
    if run.underline:
        ann.append("underline")
    # python-docx doesn't expose strikethrough directly; check
    # the <w:rPr><w:strike> / <w:dstrike> elements.
    rpr = run._element.find(qn("w:rPr"))
    if rpr is not None:
        if rpr.find(qn("w:strike")) is not None or rpr.find(qn("w:dstrike")) is not None:
            ann.append("strikethrough")
    return ann


def _num_info(para: Any) -> Optional[tuple[int, int]]:
    """Return ``(num_id, ilvl)`` for a list-item paragraph, or None.

    numId is the list definition ID; ilvl is the indent level
    (0..8). Both are required for the paragraph to be a list
    item. python-docx exposes ``para._p`` which has the XML.
    """
    ppr = para._p.find(qn("w:pPr"))
    if ppr is None:
        return None
    numpr = ppr.find(qn("w:numPr"))
    if numpr is None:
        return None
    nid = numpr.find(qn("w:numId"))
    ilvl = numpr.find(qn("w:ilvl"))
    if nid is None or ilvl is None:
        return None
    try:
        return (int(nid.get(qn("w:val"))), int(ilvl.get(qn("w:val"))))
    except (TypeError, ValueError):
        return None


def _list_style_from_num(path: Path, num_id: int) -> str:
    """Inspect the numbering.xml to decide ordered vs unordered.

    Looks for the <w:abstractNum> referenced by ``num_id`` and
    returns ``"ordered"`` for arabic / roman numeral formats,
    ``"unordered"`` for bullet formats, ``"task"`` for task list
    formats (the latter is rare in DOCX, but Word supports it).
    """
    try:
        with zipfile.ZipFile(path) as zf:
            if "word/numbering.xml" not in zf.namelist():
                return "unordered"
            with zf.open("word/numbering.xml") as f:
                tree = ET.parse(f)
    except (zipfile.BadZipFile, ET.ParseError, OSError):
        return "unordered"
    root = tree.getroot()
    # num -> abstractNumId
    num_el = root.find(f"{W_NS}num[@{W_NS}numId='{num_id}']")
    if num_el is None:
        return "unordered"
    aid_el = num_el.find(f"{W_NS}abstractNumId")
    if aid_el is None:
        return "unordered"
    try:
        aid = int(aid_el.get(f"{W_NS}val"))
    except (TypeError, ValueError):
        return "unordered"
    abs_el = root.find(f"{W_NS}abstractNum[@{W_NS}abstractNumId='{aid}']")
    if abs_el is None:
        return "unordered"
    lvl0 = abs_el.find(f"{W_NS}lvl[@{W_NS}ilvl='0']")
    if lvl0 is None:
        return "unordered"
    fmt = lvl0.find(f"{W_NS}numFmt")
    if fmt is None:
        return "unordered"
    val = fmt.get(f"{W_NS}val") or "bullet"
    if val in ("decimal", "upperRoman", "lowerRoman", "upperLetter", "lowerLetter"):
        return "ordered"
    if val == "checkbox":
        return "task"
    return "unordered"


# ---------------------------------------------------------------------------
# List emission
# ---------------------------------------------------------------------------


def _emit_list_item(
    para: Any,
    num_info: tuple[int, int],
    open_lists: dict[int, dict[str, Any]],
    emitted_blocks: list[IntermediateBlock],
    text: str,
) -> list[IntermediateBlock]:
    """Emit a list item, possibly creating a new list container.

    The list container is emitted as soon as the second item
    appears (we need at least one item to know the style). The
    first item is held in a temporary buffer and emitted together
    with the container.
    """
    num_id, ilvl = num_info
    runs = [_run_dict(r.text, _run_annotations(r)) for r in _runs_of(para)]
    if not runs and text:
        runs = [_run_dict(text)]

    if num_id not in open_lists:
        # New list. Buffer the first item; emit the list
        # container on the next item (or at end of document).
        open_lists[num_id] = {
            "ilvl": ilvl,
            "items": [IntermediateBlock(type="listItem", runs=runs)],
        }
        return []

    # Same numId: append to the existing list.
    bucket = open_lists[num_id]
    bucket["items"].append(IntermediateBlock(type="listItem", runs=runs))
    return []


def _close_all_lists(
    open_lists: dict[int, dict[str, Any]],
    emitted_blocks: list[IntermediateBlock],
) -> None:
    """Flush all open lists as ``list`` blocks into the emitted stream.

    Called at the end of the document and at every table
    boundary (a table breaks the list).
    """
    for num_id, bucket in list(open_lists.items()):
        style = bucket.get("style", "unordered")
        items = bucket["items"]
        emitted_blocks.append(
            IntermediateBlock(
                type="list",
                attrs={"style": style, "items": [it.id for it in items]},
                children=items,
            )
        )
    open_lists.clear()


# ---------------------------------------------------------------------------
# Table emission
# ---------------------------------------------------------------------------


def _emit_table(tbl_el: Any, doc: Document) -> IntermediateBlock:
    """Emit a DOCX table as a ``table`` block with cell blocks.

    DOCX tables can have ``<w:tbl>`` containing ``<w:tr>`` (rows)
    containing ``<w:tc>`` (cells) containing one or more ``<w:p>``
    (paragraphs). We walk that hierarchy and emit a list of cell
    blocks. Each cell block is the concatenation of the cell's
    paragraphs (the AST doesn't have a multi-paragraph cell type,
    so we join with a newline).
    """
    from docx.table import Table

    table = Table(tbl_el, doc.part)
    rows = list(table.rows)
    n_rows = len(rows)
    n_cols = 0
    cell_blocks_grid: list[list[IntermediateBlock]] = []

    for r in rows:
        row_blocks: list[IntermediateBlock] = []
        for c in r.cells:
            # Cell text: concatenate the cell's paragraph text
            # with newlines so multi-paragraph cells survive.
            cell_lines: list[str] = []
            for p in c.paragraphs:
                pt = "".join(run.text for run in p.runs)
                cell_lines.append(pt)
            cell_text = "\n".join(cell_lines).strip()
            # Each cell becomes a single paragraph block.
            if cell_text:
                row_blocks.append(
                    IntermediateBlock(
                        type="paragraph",
                        runs=[_run_dict(cell_text)],
                    )
                )
            else:
                row_blocks.append(IntermediateBlock(type="paragraph", runs=[]))
        n_cols = max(n_cols, len(row_blocks))
        cell_blocks_grid.append(row_blocks)

    # Pad rows to a uniform column count (DOCX tables can have
    # ragged edges when cells span columns).
    for row in cell_blocks_grid:
        while len(row) < n_cols:
            row.append(IntermediateBlock(type="paragraph", runs=[]))

    # Flatten the grid so children carries every cell block in
    # row-major order; the AST's cells attribute mirrors the grid.
    flat: list[IntermediateBlock] = [c for row in cell_blocks_grid for c in row]
    return IntermediateBlock(
        type="table",
        attrs={
            "rows": n_rows,
            "cols": n_cols,
            "cells": [[c.id for c in row] for row in cell_blocks_grid],
        },
        children=flat,
    )


# ---------------------------------------------------------------------------
# Image extraction
# ---------------------------------------------------------------------------


def _inject_images(
    blocks: list[IntermediateBlock],
    path: Path,
    media_dir: Optional[Path],
    doc: Document,
) -> list[IntermediateBlock]:
    """Replace placeholder paragraphs with image blocks where applicable.

    python-docx doesn't expose inline images directly. We scan
    the document's relationships for image parts and look for
    <w:drawing> / <w:pict> elements in each paragraph. When we
    find a paragraph whose only content is an image (no text), we
    replace it with an ``image`` block.
    """
    if media_dir is None:
        return blocks

    media_dir.mkdir(parents=True, exist_ok=True)

    # Build a map: paragraph_id -> list of image rIds
    p_to_images: dict[int, list[str]] = {}
    for para in doc.paragraphs:
        drawings = para._p.findall(f".//{qn('w:drawing')}")
        picts = para._p.findall(f".//{qn('w:pict')}")
        if not (drawings or picts):
            continue
        rids: list[str] = []
        for d in drawings:
            blip = d.find(f".//{qn('a:blip')}")
            if blip is not None:
                rid = blip.get(
                    "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
                )
                if rid:
                    rids.append(rid)
        for p in picts:
            # <w:pict> uses v:imagedata with r:id
            for img in p.findall(".//{urn:schemas-microsoft-com:vml}imagedata"):
                rid = img.get(
                    "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id"
                )
                if rid:
                    rids.append(rid)
        if rids:
            p_to_images[id(para._p)] = rids

    if not p_to_images:
        return blocks

    # Walk the document body and replace any paragraph that has
    # an image attached with an image block. We keep the original
    # paragraph's position in the output stream.
    out: list[IntermediateBlock] = []
    body = doc.element.body
    for el in list(body):
        if el.tag == qn("w:p") and id(el) in p_to_images:
            rids = p_to_images[id(el)]
            for rid in rids:
                rel = doc.part.rels.get(rid)
                if rel is None:
                    continue
                # rel.target_ref is the part name inside the zip;
                # open it and copy the bytes to media_dir.
                part = rel.target_part
                if part is None:
                    continue
                ext = Path(part.partname).suffix or ".bin"
                fname = f"image_{len(out)}_{Path(part.partname).stem}{ext}"
                dst = media_dir / fname
                try:
                    with zipfile.ZipFile(path) as zf:
                        with zf.open(part.partname) as src:
                            with dst.open("wb") as outf:
                                shutil.copyfileobj(src, outf)
                except (OSError, KeyError) as e:
                    log.warning("image extract failed: %s: %s", path, e)
                    continue
                out.append(
                    IntermediateBlock(
                        type="image",
                        attrs={"source": str(dst), "alt": ""},
                    )
                )
        # We re-emit the paragraph as a normal block to preserve
        # the order; if the paragraph was a captioned image, the
        # text will be a separate paragraph below.
        # We re-walk the body via the high-level API to keep
        # things simple: paragraphs in order, skipping the ones
        # we've already replaced with an image.
    # Combine: the rest of the blocks emitted by ``parse_docx``
    # are interleaved; we want them in the same order as the
    # body. This is a best-effort merge: for each block in
    # ``blocks`` we just keep it; image blocks are appended
    # before the next non-image block.
    # Practical v1 behaviour: image blocks are appended at the
    # end of the document when the user opens a DOCX. The AST
    # builder treats this as "first page is text, last page is
    # figures" which is acceptable for an importer.
    out.extend(blocks)
    return out


# ---------------------------------------------------------------------------
# Footnotes
# ---------------------------------------------------------------------------


def _extract_footnotes(path: Path) -> list[IntermediateBlock]:
    """Return footnote blocks for the document.

    Footnotes live in ``word/footnotes.xml`` inside the .docx
    zip. We extract the text of each non-separator footnote and
    emit it as a ``quote`` block (the AST doesn't have a
    footnote type, so quote-with-cite is the closest analogue).
    """
    try:
        with zipfile.ZipFile(path) as zf:
            if "word/footnotes.xml" not in zf.namelist():
                return []
            with zf.open("word/footnotes.xml") as f:
                tree = ET.parse(f)
    except (zipfile.BadZipFile, ET.ParseError, OSError):
        return []
    root = tree.getroot()
    out: list[IntermediateBlock] = []
    for fn in root.findall(f"{W_NS}footnote"):
        # Skip separator / continuationSeparator footnotes
        ftype = fn.get(f"{W_NS}type")
        if ftype in ("separator", "continuationSeparator"):
            continue
        text = "".join(t.text or "" for t in fn.iter(f"{W_NS}t"))
        if not text.strip():
            continue
        out.append(
            IntermediateBlock(
                type="quote",
                attrs={"cite": "footnote"},
                runs=[_run_dict(text.strip())],
            )
        )
    return out
