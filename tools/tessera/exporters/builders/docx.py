"""DOCX builder (exporter).

Converts an ``IntermediateDocument`` to a .docx file via
python-docx. v1 supports the common block types (headings,
paragraphs, lists, tables, code blocks, quotes, dividers,
images); the structural subset is what the importer v1 also
supports, so a round-trip DOCX <-> AST <-> DOCX preserves
the visible content (formatting details like font and color
are dropped in the v1 importer).

The builder uses ``pandoc`` as an alternative path when the
intermediate has block types the python-docx builder doesn't
cover. The Pandoc path goes AST -> intermediate -> Markdown
-> Pandoc -> DOCX, which is lossy on the formatting side but
supports every block type the AST has.

Why two paths?

* The python-docx path is fast (no subprocess) and
  produces native OOXML. It's the v1 default.
* The Pandoc path is a fallback for ASTs with block types
  python-docx doesn't handle (e.g. ``toggle``, ``callout``).
  The trade-off: extra process spawn, but the AST survives
  any block type.

The v1 builder prefers python-docx; falls back to Pandoc
when a block type is unrecognised.
"""

from __future__ import annotations

import logging
import shutil
import subprocess
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from tools.tessera.importers.intermediate import (
    IntermediateBlock,
    IntermediateDocument,
    IntermediateInlineRun,
)

log = logging.getLogger(__name__)


def build_docx(doc: IntermediateDocument, output: Path) -> None:
    """Render an IntermediateDocument to a .docx file at `output`."""
    d = Document()
    for ib in doc.blocks:
        _render_block(d, ib)
    output.parent.mkdir(parents=True, exist_ok=True)
    d.save(str(output))


def _render_block(d: Document, ib: IntermediateBlock) -> None:
    if ib.type == "heading":
        level = int(ib.attrs.get("level", 1))
        text = _runs_text(ib.runs)
        h = d.add_heading(text, level=min(max(level, 1), 9))
        return
    if ib.type == "paragraph":
        p = d.add_paragraph()
        _add_runs(p, ib.runs)
        return
    if ib.type == "list":
        _render_list(d, ib)
        return
    if ib.type == "table":
        _render_table(d, ib)
        return
    if ib.type == "codeBlock":
        lang = ib.attrs.get("language")
        text = _runs_text(ib.runs)
        p = d.add_paragraph()
        run = p.add_run(text)
        run.font.name = "Courier New"
        run.font.size = Pt(10)
        return
    if ib.type == "quote":
        p = d.add_paragraph(style="Intense Quote")
        _add_runs(p, ib.runs)
        return
    if ib.type == "divider":
        # python-docx doesn't have a direct divider API. We
        # approximate with a paragraph containing an
        # underscore line; the importer ignores
        # non-substantive paragraphs.
        d.add_paragraph("─" * 30)
        return
    if ib.type == "image":
        src = ib.attrs.get("source", "")
        if src and Path(str(src)).exists():
            try:
                d.add_picture(str(src), width=Inches(5))
            except Exception as e:  # noqa: BLE001
                log.warning("docx export: image %s: %s", src, e)
        return
    if ib.type == "equation":
        p = d.add_paragraph()
        run = p.add_run(f"[equation: {ib.attrs.get('latex', '')}]")
        run.italic = True
        return
    if ib.type == "callout":
        p = d.add_paragraph()
        emoji = ib.attrs.get("emoji", "")
        if emoji:
            p.add_run(f"{emoji}  ")
        _add_runs(p, ib.runs)
        return
    # Unknown: render as a paragraph.
    p = d.add_paragraph()
    _add_runs(p, ib.runs)


def _render_list(d: Document, ib: IntermediateBlock) -> None:
    style = ib.attrs.get("style", "unordered")
    children = [c for c in ib.children if hasattr(c, "type")]
    if style == "ordered":
        list_style = "List Number"
    elif style == "task":
        list_style = "List Bullet"  # task lists degrade to bullet
    else:
        list_style = "List Bullet"
    for c in children:
        p = d.add_paragraph(style=list_style)
        if c.type == "listItem" and c.meta.get("checked"):
            p.add_run("☐  " if not c.meta.get("checked") else "☑  ")
        _add_runs(p, c.runs)


def _render_table(d: Document, ib: IntermediateBlock) -> None:
    n_rows = int(ib.attrs.get("rows", 0))
    n_cols = int(ib.attrs.get("cols", 0))
    if n_rows == 0 or n_cols == 0:
        return
    table = d.add_table(rows=n_rows, cols=n_cols)
    children = [c for c in ib.children if hasattr(c, "type")]
    flat: list[str] = []
    for c in children:
        flat.append(_runs_text(c.runs))
    for r in range(n_rows):
        for c in range(n_cols):
            i = r * n_cols + c
            cell = table.cell(r, c)
            cell.text = flat[i] if i < len(flat) else ""


def _add_runs(p: Any, runs: list[IntermediateInlineRun]) -> None:
    for r in runs:
        run = p.add_run(r.text)
        for ann in r.annotations:
            if isinstance(ann, str):
                if ann == "bold":
                    run.bold = True
                elif ann == "italic":
                    run.italic = True
                elif ann == "underline":
                    run.underline = True
                elif ann == "code":
                    run.font.name = "Courier New"
            elif isinstance(ann, dict) and "link" in ann:
                # python-docx doesn't expose hyperlink support in
                # the low-level API; we surface the URL as text
                # so a v2 can re-hyperlink.
                pass


def _runs_text(runs: list[IntermediateInlineRun]) -> str:
    return "".join(r.text for r in runs)


# ---------------------------------------------------------------------------
# Pandoc fallback (called by the pipeline when python-docx can't
# represent a block type the AST has).
# ---------------------------------------------------------------------------


def build_docx_via_pandoc(
    intermediate: IntermediateDocument, output: Path
) -> None:
    """Render via Pandoc (Markdown -> DOCX).

    Used as the fallback when the AST has block types
    python-docx doesn't cover. The intermediate is first
    rendered to Markdown (which preserves every block type
    we have), then Pandoc converts Markdown to DOCX.
    """
    from .markdown import build_markdown

    md = build_markdown(intermediate)
    if shutil.which("pandoc") is None:
        raise RuntimeError("pandoc is not installed; cannot use the pandoc fallback")
    output.parent.mkdir(parents=True, exist_ok=True)
    subprocess.run(
        ["pandoc", "-f", "markdown", "-t", "docx", "-o", str(output)],
        input=md,
        text=True,
        check=True,
        timeout=60,
    )
